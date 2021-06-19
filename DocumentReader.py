from abc import ABC, abstractmethod
from docx.api import Document


class DocumentReader(ABC):

    def __init__(self, file, spec):
        self._file = file
        self._spec = spec
        self.tables = {}

    @abstractmethod
    def read_doc(self):
        pass


class PdfReader(DocumentReader):

    def read_doc(self):
        pass


class DocxReader(DocumentReader):
    """
    Returns a list of dictionaries for each row [{column:cell, column:cell},
                                                 {column:cell, column:cell}]
    """
    def __init__(self, file, spec):
        super().__init__(file, spec)
        self.read_doc()
    def read_doc(self):
        """returns data as rows"""
        document = Document(self._file)
        
        data = []
        for table in document.tables:
            for i, row in enumerate(table.rows):
                text = [cell.text for cell in row.cells]
                if i == 0:
                    keys = text
                    continue
                row_data = dict(zip(keys, text))
                data.append(row_data)

        # filter rows to tables we specified
        for i, spec in enumerate(self._spec):
            self.tables[str(i)] = [row for row in data if row.keys() == set(spec)]

    